from __future__ import annotations

import argparse
import html
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
REPORTS_DIR = PROJECT_ROOT / "reports"
DEFAULT_OUTPUT_DIR = PROJECT_ROOT / "outputs"


def main() -> None:
    parser = argparse.ArgumentParser(description="Export the latest Markdown report to HTML and Word.")
    parser.add_argument("--report", type=Path, default=None, help="Markdown report path. Defaults to latest report.")
    parser.add_argument("--output-dir", type=Path, default=DEFAULT_OUTPUT_DIR, help="Directory for HTML and DOCX.")
    args = parser.parse_args()

    report = args.report or latest_report()
    output_dir = args.output_dir
    output_dir.mkdir(parents=True, exist_ok=True)

    markdown = report.read_text(encoding="utf-8")
    date_part = report.stem.replace("daily-ai-brief-", "")[:8]
    html_out = output_dir / f"daily-ai-brief-{date_part}.html"
    docx_out = output_dir / f"daily-ai-brief-{date_part}.docx"

    build_html(markdown, html_out)
    build_docx(markdown, docx_out)

    print(html_out)
    print(docx_out)


def latest_report() -> Path:
    reports = sorted(REPORTS_DIR.glob("daily-ai-brief-*.md"))
    if not reports:
        raise FileNotFoundError("No reports found. Run `python3 -m src.main --config config.example.json` first.")
    return reports[-1]


def build_html(markdown: str, path: Path) -> None:
    body = markdown_to_html(markdown)
    css = """
    :root {
      color-scheme: light;
      --ink: #182230;
      --muted: #667085;
      --line: #d0d5dd;
      --soft: #f8fafc;
      --accent: #2563eb;
      --accent-soft: #eff6ff;
      --green-soft: #ecfdf3;
    }
    * { box-sizing: border-box; }
    body {
      margin: 0;
      font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", "PingFang SC", "Hiragino Sans GB", "Microsoft YaHei", Arial, sans-serif;
      color: var(--ink);
      background: #f3f4f6;
      line-height: 1.62;
    }
    main {
      max-width: 980px;
      margin: 32px auto;
      padding: 42px 52px;
      background: #fff;
      border: 1px solid #eaecf0;
      box-shadow: 0 18px 48px rgba(15, 23, 42, 0.08);
    }
    h1 {
      margin: 0 0 18px;
      font-size: 32px;
      line-height: 1.2;
      letter-spacing: 0;
      color: #101828;
    }
    h2 {
      margin: 34px 0 12px;
      padding-top: 18px;
      border-top: 1px solid var(--line);
      font-size: 22px;
      color: #101828;
    }
    h3 {
      margin: 28px 0 12px;
      padding: 18px 20px 14px;
      background: var(--accent-soft);
      border-left: 4px solid var(--accent);
      font-size: 18px;
      border-radius: 6px;
    }
    h4 {
      margin: 18px 0 8px;
      font-size: 15px;
      color: #344054;
    }
    p { margin: 8px 0 14px; }
    ul { margin: 8px 0 16px 22px; padding: 0; }
    li { margin: 5px 0; }
    a { color: #175cd3; text-decoration: none; }
    a:hover { text-decoration: underline; }
    code, pre {
      font-family: ui-monospace, SFMono-Regular, Menlo, Consolas, monospace;
    }
    pre {
      margin: 10px 0 16px;
      padding: 14px 16px;
      background: #101828;
      color: #f9fafb;
      border-radius: 6px;
      overflow-x: auto;
      white-space: pre-wrap;
    }
    .note {
      padding: 12px 16px;
      background: var(--green-soft);
      border: 1px solid #abefc6;
      border-radius: 6px;
    }
    @media (max-width: 720px) {
      main { margin: 0; padding: 24px 18px; border: 0; box-shadow: none; }
      h1 { font-size: 26px; }
      h2 { font-size: 20px; }
    }
    """
    mermaid_script = """
    <script type="module">
      import mermaid from 'https://cdn.jsdelivr.net/npm/mermaid@11/dist/mermaid.esm.min.mjs';
      mermaid.initialize({ startOnLoad: true, theme: 'default' });
    </script>
    """
    path.write_text(
        "<!doctype html>\n"
        "<html lang=\"zh-CN\"><head><meta charset=\"utf-8\">"
        "<meta name=\"viewport\" content=\"width=device-width, initial-scale=1\">"
        f"<title>每日 AI 研究简报</title><style>{css}</style></head>"
        f"<body><main>{body}</main></body>{mermaid_script}</html>",
        encoding="utf-8",
    )


def markdown_to_html(markdown: str) -> str:
    lines = markdown.splitlines()
    out: list[str] = []
    in_list = False
    in_code = False
    code_type = ""
    code_lines: list[str] = []
    for raw in lines:
        line = raw.rstrip()
        if line.startswith("```"):
            if in_code:
                code = html.escape("\n".join(code_lines))
                if code_type == "mermaid":
                    out.append(f'<pre class="mermaid">{code}</pre>')
                else:
                    out.append("<pre><code>" + code + "</code></pre>")
                code_lines = []
                code_type = ""
                in_code = False
            else:
                if in_list:
                    out.append("</ul>")
                    in_list = False
                code_type = line.replace("```", "").strip()
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if not line:
            if in_list:
                out.append("</ul>")
                in_list = False
            continue
        if line.startswith("# "):
            if in_list:
                out.append("</ul>")
                in_list = False
            out.append(f"<h1>{inline(line[2:])}</h1>")
        elif line.startswith("## "):
            if in_list:
                out.append("</ul>")
                in_list = False
            out.append(f"<h2>{inline(line[3:])}</h2>")
        elif line.startswith("### "):
            if in_list:
                out.append("</ul>")
                in_list = False
            out.append(f"<h3>{inline(line[4:])}</h3>")
        elif line.startswith("#### "):
            if in_list:
                out.append("</ul>")
                in_list = False
            out.append(f"<h4>{inline(line[5:])}</h4>")
        elif line.startswith("- "):
            if not in_list:
                out.append("<ul>")
                in_list = True
            out.append(f"<li>{inline(line[2:])}</li>")
        else:
            if in_list:
                out.append("</ul>")
                in_list = False
            cls = " class=\"note\"" if "严格基于" in line else ""
            out.append(f"<p{cls}>{inline(line)}</p>")
    if in_list:
        out.append("</ul>")
    return "\n".join(out)


def inline(text: str) -> str:
    escaped = html.escape(text)
    url_re = re.compile(r"(https?://[^\s<]+)")
    return url_re.sub(lambda m: f'<a href="{m.group(1)}">{m.group(1)}</a>', escaped)


def build_docx(markdown: str, path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(15)
    styles["Heading 3"].font.name = "Arial"
    styles["Heading 3"].font.size = Pt(13)

    in_code = False
    code_lines: list[str] = []
    for raw in markdown.splitlines():
        line = raw.rstrip()
        if line.startswith("```"):
            if in_code:
                para = doc.add_paragraph()
                run = para.add_run("\n".join(code_lines))
                run.font.name = "Menlo"
                run.font.size = Pt(8.5)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(strip_markdown(line[2:]), level=1)
        elif line.startswith("## "):
            doc.add_heading(strip_markdown(line[3:]), level=2)
        elif line.startswith("### "):
            doc.add_heading(strip_markdown(line[4:]), level=3)
        elif line.startswith("#### "):
            para = doc.add_paragraph()
            run = para.add_run(strip_markdown(line[5:]))
            run.bold = True
        elif line.startswith("- "):
            doc.add_paragraph(strip_markdown(line[2:]), style="List Bullet")
        else:
            doc.add_paragraph(strip_markdown(line))
    doc.save(path)


def strip_markdown(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    return text


if __name__ == "__main__":
    main()
